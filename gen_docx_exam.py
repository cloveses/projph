import os
import math
import io
from docx import Document
from docx.shared import Pt , Inches
from docx.oxml.ns import qn
from pystrich.code128 import Code128Encoder

from ph_models import *

TEST_DATA = ("准考号:  ","姓　名:  ","性　别:  ","考　点:  ","报名点:  ")
STUD_DATA = ('18251402',"李文娟",'男','泗县一中','01中学')

def confirm_path(path):
    if not os.path.exists(path):
        os.makedirs(path)


def chg_font(obj,fontname='微软雅黑',size=None):
    ## 设置字体函数
    obj.font.name = fontname
    obj._element.rPr.rFonts.set(qn('w:eastAsia'),fontname)
    if size and isinstance(size,Pt):
        obj.font.size = size

def gen_barcode(code='123456'):
    ## 生成二维码
    f = io.BytesIO()
    Code128Encoder(code).save(f)
    return f

def init_doc(doc):
    ## 设置页边距
    distance = Inches(0.3)
    sec = doc.sections[0]
    sec.left_margin = distance
    sec.right_margin = distance
    sec.top_margin = distance
    sec.bottom_margin = distance
    ##设置默认字体
    chg_font(doc.styles['Normal'],fontname='宋体')


def one_page(doc,studs):

    layout_tab = doc.add_table(rows=4,cols=4)

    for cur_cell in (layout_tab.cell(0,0),layout_tab.cell(0,2),
            layout_tab.cell(1,0),layout_tab.cell(1,2),
            layout_tab.cell(2,0),layout_tab.cell(2,2),
            layout_tab.cell(3,0),layout_tab.cell(3,2),):
        ph = cur_cell.paragraphs[0]

        htitle = ph.add_run('2018年初中学业水平考试')
        htitle.add_break()
        title = ph.add_run('　准　考　证')
        title.add_break()
        chg_font(title,size=Pt(16))

        for left,right in zip(TEST_DATA,STUD_DATA):
            ph.add_run(left)
            run = ph.add_run(right)
            chg_font(run)
            run.bold = True
            run.add_break()

    for cur_cell in (layout_tab.cell(0,1),layout_tab.cell(0,3),
            layout_tab.cell(1,1),layout_tab.cell(1,3),
            layout_tab.cell(2,1),layout_tab.cell(2,3),
            layout_tab.cell(3,1),layout_tab.cell(3,3),):
        ph = cur_cell.paragraphs[0]
        run = ph.add_run()
        run.add_picture('aa.jpg',width=Inches(1.2))
        run.add_break()
        f = gen_barcode()
        run.add_picture(f,width=Inches(1.2))

def gen_unit_docx(dir_name,sch_name,studs,page_num=8):
    confirm_path(dir_name)
    path = os.path.join(dir_name,sch_name + '.docx')
    canv = canvas.Canvas(path,pagesize=(ID_SIZE[0]*mm,ID_SIZE[1]*mm))
    pages = math.ceil(len(studs)/page_num)

    doc = Document()
    init_doc(doc)

    for i in range(pages):
        one_page(doc,studs[i*page_num:(i+1)*page_num])
        doc.add_page_break()

    doc.save(path)

# 按学校生成准考证
@db_session
def gen_examid_sch(dir_name):
    schs = select(s.sch for s in StudPh)
    for sch in schs:
        datas = select(s 
         for s in StudPh if s.sch==sch).order_by(StudPh.classcode,StudPh.phid)
        datas = [(s.phid,s.name,s.sex,s.exam_addr,s.sch,s.schcode,s.signid) for s in datas]
        gen_unit_docx(dir_name,sch,datas)


if __name__ == '__main__':
    pass